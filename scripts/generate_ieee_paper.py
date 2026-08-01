"""
IEEE-format paper matching format.pdf exactly.
ONE section, TWO columns throughout. Title + authors + abstract + index
terms in borderless full-width tables (span columns). Body flows in 2 cols.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FIG_DIR = 'C:/Users/soumy/OneDrive/Desktop/Project/lca_platform/scripts/figures'
OUTPUT = 'C:/Users/soumy/AppData/Local/Temp/LCA_Platform_IEEE_Paper.docx'

doc = Document()

# ============ PAGE SETUP ============
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
# Two columns
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
for ec in section._sectPr.findall(qn('w:cols')):
    section._sectPr.remove(ec)
section._sectPr.append(parse_xml('<w:cols {} w:num="2" w:space="360" w:equalWidth="1"/>'.format(nsdecls('w'))))

# ============ STYLES ============
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(3)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(9)
    hs.paragraph_format.space_after = Pt(3)
    hs.paragraph_format.line_spacing = 1.0
    hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hs.font.size = Pt(11) if level == 1 else Pt(10)


def add_para(text, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=3, space_before=3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_fig(filename, width=4.8, caption=''):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        add_para(f'[Figure: {filename} not found]', italic=True, size=9)
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run()
    r.add_picture(path, width=Inches(width))
    if caption:
        add_para(caption, italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)


def add_table_ieee(headers, rows, caption=''):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'
    add_para('', size=2, space_after=2)


def remove_all_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'none')
        element.set(qn('w:sz'), '0')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        borders.append(element)
    tblPr.append(borders)


def add_authors_side_by_side():
    """Side-by-side authors using borderless 2-col table."""
    tbl = doc.add_table(rows=1, cols=2)
    remove_all_borders(tbl)
    tbl.autofit = True
    data = [
        ('Soumya Subhra Datta',
         'Dept. of Computer Science and Engineering -',
         'Artificial Intelligence',
         'Sathyabama Institute of Science',
         'and Technology',
         'Chennai, India'),
        ('Yeluri Nanda Gopal',
         'Dept. of Computer Science and Engineering -',
         'Artificial Intelligence',
         'Sathyabama Institute of Science',
         'and Technology',
         'Chennai, India'),
    ]
    for ci, (name, d1, d2, c1, c2, city) in enumerate(data):
        cell = tbl.cell(0, ci)
        cell.paragraphs[0].clear()
        lines = [(name, True, 10), (d1, False, 9), (d2, False, 9),
                 (c1, False, 9), (c2, False, 9), (city, False, 9)]
        for ti, (txt, bld, sz) in enumerate(lines):
            p = cell.paragraphs[0] if ti == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(txt)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(sz)
            r.bold = bld
            r.italic = not bld





def add_fullwidth_cell(cell, runs_data):
    """Add runs with formatting to a single-cell table cell. runs_data: list of (text, bold, italic, size)."""
    for i, (txt, bld, ita, sz) in enumerate(runs_data):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(sz)
        r.bold = bld
        r.italic = ita


def make_fullwidth_table():
    """Create a borderless 1x1 table spanning full page width."""
    tbl = doc.add_table(rows=1, cols=1)
    remove_all_borders(tbl)
    tbl.autofit = True
    return tbl


# ============================================================
# TITLE (full-width table spanning both columns)
# ============================================================
tbl = make_fullwidth_table()
add_fullwidth_cell(tbl.cell(0, 0), [
    ('AI-Driven Life Cycle Assessment Platform for Sustainable Metal Production', True, False, 16),
])
# extra spacing below title
tbl.cell(0, 0).add_paragraph()
tbl.cell(0, 0).add_paragraph()

# ============================================================
# AUTHORS (side by side in full-width 1x2 table)
# ============================================================
add_authors_side_by_side()

# ============================================================
# ABSTRACT (full-width table)
# ============================================================
tbl2 = make_fullwidth_table()
c = tbl2.cell(0, 0)
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(3)
r1 = p.add_run('Abstract')
r1.bold = True
r1.font.size = Pt(10)
r1.font.name = 'Times New Roman'
r2 = p.add_run(
    '\u2014This paper presents an AI-driven Life Cycle Assessment (LCA) platform '
    'for evaluating environmental sustainability of metal production across 14 ore types. '
    'The platform integrates four machine learning models\u2014Gradient Boosting Regressor for '
    'Heavy Rare Earth Element (HREE) prediction (R\u00b2=0.946), Random Forest Classifier for '
    'deposit classification (Accuracy=0.769), Random Forest Regressor for resource estimation '
    '(R\u00b2=0.991), and Gradient Boosting Regressor for Dy\u2082O\u2083 content prediction '
    '(R\u00b2=0.849)\u2014trained on 15 mining datasets comprising 50,000+ records. The LCA engine '
    'models five environmental categories across the mining-to-metal lifecycle with ore-specific '
    'multipliers calibrated to industry benchmarks. A circular economy module quantifies material '
    'recovery, waste diversion, and resource efficiency metrics. The multi-dimensional sustainability '
    'scoring system evaluates Environmental, Social, Governance, Economic, and Innovation (ESG+E) '
    'performance. SHAP-based model explainability provides transparent, interpretable predictions. '
    'The platform is deployed as a production-ready full-stack web application using FastAPI and Docker.')
r2.font.size = Pt(9)
r2.font.name = 'Times New Roman'

# ============================================================
# INDEX TERMS (full-width table)
# ============================================================
tbl3 = make_fullwidth_table()
c = tbl3.cell(0, 0)
p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r1 = p.add_run('Index Terms')
r1.bold = True
r1.font.size = Pt(9)
r1.font.name = 'Times New Roman'
r2 = p.add_run(
    '\u2014Life Cycle Assessment, Machine Learning, Gradient Boosting, '
    'Random Forest, Rare Earth Elements, Sustainability, Circular Economy, SHAP')
r2.font.size = Pt(9)
r2.font.name = 'Times New Roman'

# ============================================================
# I. INTRODUCTION
# ============================================================
doc.add_heading('I. Introduction', level=1)
add_para(
    'The extraction and processing of metals\u2014particularly Rare Earth Elements (REEs)\u2014'
    'represent one of the most environmentally consequential industrial activities globally. '
    'The production of one tonne of REE oxides generates approximately 2,000 tonnes of toxic waste, '
    'including radioactive residues, acid mine drainage, and significant CO\u2082 emissions. '
    'As the global transition to clean energy accelerates, demand for critical minerals including '
    'lithium, cobalt, nickel, and rare earths is projected to increase 4-6 fold by 2040 [1], '
    'making rigorous environmental assessment an industrial and ecological imperative.')

add_para(
    'Traditional Life Cycle Assessment (LCA) methodologies, governed by ISO 14040/14044 standards, '
    'provide systematic frameworks for evaluating environmental impacts across a product\'s lifecycle. '
    'However, conventional LCA suffers from significant limitations: data collection is labor-intensive, '
    'emission factors are often outdated, cross-commodity comparisons are difficult, and the integration '
    'of circular economy principles remains nascent [2], [3].')

add_para(
    'This paper introduces a machine learning-augmented LCA platform that addresses these limitations '
    'through four key contributions: (1) ML-driven prediction of geological and compositional parameters '
    'with R\u00b2 up to 0.991, (2) deterministic LCA modelling across 14 metal types with ore-specific '
    'environmental multipliers, (3) circular economy quantification with industry benchmarking across '
    'five circularity metrics, and (4) multi-dimensional sustainability scoring aligned with ESG '
    'frameworks enhanced by SHAP-based explainability.')

# ============================================================
# II. LITERATURE REVIEW
# ============================================================
doc.add_heading('II. Literature Review', level=1)

add_para(
    'Life Cycle Assessment has been applied to mining operations extensively in the literature. '
    'Huijbregts [4] established midpoint impact categories for mineral extraction, while Norgate '
    'and Jahanshahi [5] developed emission factors for base metal smelting operations. '
    'More recently, the European Commission Critical Raw Materials Act [6] identified 34 critical '
    'materials requiring comprehensive circular economy strategies.')

add_para(
    'Machine learning applications in mineral exploration have grown substantially. '
    'Zuo et al. [7] demonstrated Random Forest classifiers achieving 78% accuracy in deposit type '
    'prediction using geochemical features. Rodriguez-Galiano et al. [8] applied Gradient Boosting '
    'to geochemical anomaly detection with superior performance compared to traditional logistic '
    'regression. Carranza [9] provided foundational work on data-driven mineral prospectivity mapping.')

add_para(
    'Recent work by Liu et al. [10] employed SHAP (SHapley Additive exPlanations) for interpretable '
    'mineral prospectivity mapping, demonstrating that model transparency is critical for geological '
    'acceptance. Lundberg and Lee [11] originally proposed SHAP as a unified framework for interpreting '
    'predictions based on Shapley values from cooperative game theory.')

add_para(
    'Circular economy frameworks for metals have been developed by Linder et al. [12], who established '
    'methodologies for measuring product-level circularity. Kumar and Singh [13] provided multi-criteria '
    'decision-making frameworks for sustainable mining operations, while Chen et al. [14] examined '
    'life cycle impacts of lithium-ion battery recycling.')

# ============================================================
# III. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('III. System Architecture', level=1)

add_para(
    'The platform follows a modular, service-oriented architecture with five core engines: '
    'Data Pipeline Engine (loads and preprocesses 15 datasets from three primary sources), '
    'Machine Learning Engine (four trained models with StandardScaler normalization and 5-fold '
    'cross-validation), Life Cycle Assessment Engine (deterministic calculators for five impact '
    'categories across 14 ore types), Circular Economy Engine (recycling, waste diversion, and '
    'resource efficiency scoring), and Sustainability Scoring Engine (ESG+E scoring with industry '
    'benchmarking).')

add_para(
    'The backend is implemented in Python 3.11 using FastAPI for RESTful API routing with '
    'asynchronous endpoint handling, SQLAlchemy for ORM-based database management with SQLite, '
    'and scikit-learn 1.3.2 for ML model training and inference. The frontend is a single-page '
    'application using vanilla JavaScript with Chart.js for interactive data visualization '
    'and JWT-based authentication using PBKDF2 password hashing.')

add_fig('fig9_system_architecture.png', 5.2,
        'Fig. 1. System architecture showing the layered design from data sources through ML and LCA engines to deployment.')

# ============================================================
# IV. DATASETS AND DATA ENGINEERING
# ============================================================
doc.add_heading('IV. Datasets and Data Engineering', level=1)

add_para(
    'The platform integrates data from three primary sources, yielding 15 distinct datasets spanning '
    'global mining operations. The primary sources are: Global Rare Earth Elements Projects database '
    '(mining_projects.csv with 475+ projects and 15 REE oxide compositions), Open Database on Global '
    'Coal and Metal Mining (11 tabular datasets plus one geospatial file covering active mines, '
    'production statistics, and environmental indicators), and World Mining Commodities database '
    '(annual production across 30+ commodities).')

doc.add_heading('4.1 Feature Engineering', level=2)

add_para(
    'The preprocessing pipeline applies the following transformations: logarithmic transform '
    '(log\u2081\u2080 of resource tonnes), normalized oxide ratios (individual oxide / sum of all oxides \u00d7 100), '
    'LREE/HREE summation and ratio calculation, label encoding of categorical variables (continent, '
    'deposit type, country), interaction features (pairwise oxide products for selected element pairs), '
    'StandardScaler normalization to zero mean and unit variance, and median imputation for missing values. '
    'The final feature vector comprises log\u2081\u2080(resource), grade percentage, continent, deposit type, '
    'and 15 normalized REE oxide percentages.')

add_fig('fig10_data_workflow.png', 5.2,
        'Fig. 2. End-to-end data processing and model training workflow from raw data through feature engineering to ML models.')

# ============================================================
# V. MACHINE LEARNING MODELS
# ============================================================
doc.add_heading('V. Machine Learning Models', level=1)

add_para(
    'Four ML models were developed for geological and compositional property prediction. '
    'Each model uses 5-fold cross-validation with the dataset split 80/20 train/test '
    '(random_state=42 for reproducibility). Feature selection is performed using recursive '
    'feature elimination with cross-validation (RFECV) to identify the optimal feature subset. '
    'Performance metrics include coefficient of determination (R\u00b2), Root Mean Square Error (RMSE), '
    'classification accuracy, and F1-weighted score.')

doc.add_heading('5.1 HREE Percentage Predictor', level=2)
add_para(
    'Architecture: Gradient Boosting Regressor with n_estimators=200, max_depth=6, '
    'learning_rate=0.1, and subsample=0.8. The model predicts the Heavy Rare Earth Element (HREE) '
    'percentage using geochemical and geological features. Gradient Boosting minimizes mean squared '
    'error iteratively through additive weak learners:')
add_para(
    'F\u2098(x) = F\u2098\u208b\u2081(x) + \u03b7 \u00b7 h\u2098(x)   where   '
    'h\u2098 = argmin\u209c \u2211\u1d62 L(y\u1d62, F\u2098\u208b\u2081(x\u1d62) + h(x\u1d62))   (1)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'Performance: Test R\u00b2 = 0.946, RMSE = 7.50%, 5-fold CV R\u00b2 = 0.929. '
    'The top predictive features are Dy\u2082O\u2083 content (SHAP value: +0.42), '
    'Y\u2082O\u2083 content (SHAP value: +0.35), and deposit type (Ionic Clay: +0.28).')

add_fig('fig2_hree_features.png', 4.5,
        'Fig. 3. Feature importance for HREE Percentage Predictor showing Dy\u2082O\u2083 and Y\u2082O\u2083 as dominant features.')

doc.add_heading('5.2 Deposit Type Classifier', level=2)
add_para(
    'Architecture: Random Forest Classifier with n_estimators=200, max_depth=10, and '
    'class_weight="balanced" to handle class imbalance across six deposit types. '
    'The Random Forest combines bootstrap aggregating (bagging) with random feature '
    'selection, predicting via majority voting:')
add_para(
    '\u0237 = mode{ h(x, \u03b8\u2081), h(x, \u03b8\u2082), ..., h(x, \u03b8\u2097) }   (2)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'Performance: Accuracy = 0.769, F1 Weighted = 0.765, 5-fold CV Accuracy = 0.734. '
    'The classifier achieves highest precision for Carbonatite deposits (0.84) and '
    'Ionic Clay deposits (0.81), while confusion is most common between '
    'Alkaline Intrusive and Placer deposit types.')

doc.add_heading('5.3 Resource Size Estimator', level=2)
add_para(
    'Architecture: Random Forest Regressor with n_estimators=200 and max_depth=8. '
    'The target variable is log\u2081\u2080(resource tonnes), which stabilizes variance across '
    'the wide range of deposit sizes (from 10\u2074 to 10\u2071\u2070 tonnes). '
    'Training minimizes mean squared error on log-transformed values:')
add_para(
    '\u0177 = log\u2081\u2080(resource)   where   L(y, \u0177) = (log\u2081\u2080(y) \u2212 \u0177)\u00b2   (3)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'Performance: Test R\u00b2 = 0.991, CV R\u00b2 = 0.919. The high R\u00b2 indicates excellent '
    'fit on log-transformed data, with grade percentage and deposit type emerging as '
    'the most influential predictors.')

doc.add_heading('5.4 Dy\u2082O\u2083 Content Predictor', level=2)
add_para(
    'Architecture: Gradient Boosting Regressor with n_estimators=200, max_depth=5, and '
    'learning_rate=0.1. This model predicts Dysprosium oxide (Dy\u2082O\u2083) content, a critical '
    'heavy rare earth element essential for permanent magnets in electric vehicles and wind turbines. '
    'Dysprosium is listed as a critical raw material by the European Commission due to supply risk '
    'and economic importance [6].')
add_para(
    'Performance: Test R\u00b2 = 0.849, 5-fold CV R\u00b2 = 0.943. The high cross-validation '
    'score relative to test score reflects the model\'s consistency across different data splits. '
    'Key predictors include HREE ratio, Y\u2082O\u2083 content, and deposit type (Carbonatite).')

add_fig('fig3_dy_features.png', 4.5,
        'Fig. 4. Feature importance for Dy\u2082O\u2083 Content Predictor showing HREE ratio and Y\u2082O\u2083 as key predictors.')

# ============================================================
# VI. LIFE CYCLE ASSESSMENT ENGINE
# ============================================================
doc.add_heading('VI. Life Cycle Assessment Engine', level=1)

add_para(
    'The LCA engine implements process-based attributional LCA modelling for five environmental '
    'impact categories across the mining-to-metal lifecycle. Each category uses ore-specific '
    'multipliers calibrated to industry benchmarks from the literature [4], [5], [15]. '
    'The engine follows ISO 14040/14044 standards with four phases: goal and scope definition, '
    'life cycle inventory analysis, life cycle impact assessment, and interpretation.')

doc.add_heading('6.1 Carbon Footprint', level=2)
add_para(
    'Total CO\u2082 emissions are calculated as the sum of three lifecycle stages:')
add_para(
    'CO\u2082\u209c\u2092\u209c\u2090\u2113 = E\u2098\u1d62\u2099\u1d62\u0274\u0262 + E\u1d18\u0280\u1d57\u1d04\u1d07\u0455\u0455\u1d62\u0274\u0262 + E\u1d20\u0280\u2092\u0274\u0455\u1d18\u1d57\u0280\u1d62   (4)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'where each stage is computed as: E\u1d62 = M\u2092\u0280\u2091 \u00d7 EF\u1d62 \u00d7 M_carbon. '
    'The mining emission factor depends on extraction method: surface mining at 12.0 kg CO\u2082/t ore '
    'and underground mining at 25.0 kg CO\u2082/t ore. Carbon multipliers across ore types range from '
    '0.6 (Uranium) to 2.8 (Aluminium), reflecting processing energy intensity.')

doc.add_heading('6.2 Environmental Impact Score', level=2)
add_para(
    'The composite environmental impact score integrates five categories with weighted contributions:')
add_para(
    'EIS = (C\u2099\u2092\u0280\u2098 \u00d7 0.30 + W\u2099\u2092\u0280\u2098 \u00d7 0.20 + '
    'En\u2099\u2092\u0280\u2098 \u00d7 0.25 + Wa\u2099\u2092\u0280\u2098 \u00d7 0.15 + A\u2099\u2092\u0280\u2098 \u00d7 0.10) \u00d7 100   (5)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'Normalized category scores range from 0 (minimum impact) to 100 (maximum impact). '
    'Grade classification: A (0-15), B (16-30), C (31-50), D (51-69), E (70-100).')

add_fig('fig4_ore_multipliers.png', 5.0,
        'Fig. 5. Environmental impact multipliers across 14 ore types with Rare Earth Elements (REE) baseline normalized to 1.0.')

add_fig('fig7_lca_breakdown.png', 5.0,
        'Fig. 6. Carbon emission distribution by processing stage for REE and Copper, showing mining, processing, and transport contributions.')

add_fig('fig11_lca_workflow.png', 5.0,
        'Fig. 7. LCA assessment workflow from input parameters through five parallel environmental calculators to impact grading.')

add_table_ieee(
    ['Ore Type', 'Carbon Mult. (REE=1)', 'Water Mult. (REE=1)', 'Energy Mult. (REE=1)'],
    [['Aluminium', '2.8', '2.1', '3.2'],
     ['Copper', '1.8', '1.5', '2.0'],
     ['Iron', '1.0', '0.8', '1.2'],
     ['REE Baseline', '1.0', '1.0', '1.0'],
     ['Lithium', '2.2', '3.5', '2.8'],
     ['Uranium', '0.6', '0.8', '1.2']],
    'TABLE I. Environmental Impact Multipliers by Ore Type')

# ============================================================
# VII. CIRCULAR ECONOMY ASSESSMENT
# ============================================================
doc.add_heading('VII. Circular Economy Assessment', level=1)

add_para(
    'The circular economy module implements a multi-metric framework quantifying material circularity, '
    'waste diversion, and resource efficiency. The composite circularity score is computed as:')
add_para(
    'CS = MRR \u00d7 0.30 + WDR \u00d7 0.25 + RE \u00d7 0.20 + WRR \u00d7 0.15 + ERR \u00d7 0.10   (6)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'where MRR = (recycled material / total product output) \u00d7 100 (Material Recovery Rate), '
    'WDR = (waste diverted / total waste generated) \u00d7 100 (Waste Diversion Rate), '
    'RE = (product output / ore processed) \u00d7 100 (Resource Efficiency), '
    'WRR = (waste recycled / total waste) \u00d7 100 (Waste Recycling Rate), and '
    'ERR = (energy recovered / total energy consumed) \u00d7 100 (Energy Recovery Rate).')

add_para(
    'Industry recycling rates reveal a dramatic disparity across metal types. Iron and steel lead '
    'with approximately 90% recycling rates due to established scrap collection infrastructure and '
    'magnetic separation technology. Aluminium achieves 75% recycling with significantly lower '
    'energy requirements compared to primary production (only 5% of the energy). In contrast, '
    'rare earth elements have a recycling rate of approximately 1%, representing a critical '
    'circular economy gap driven by dispersed applications, complex product designs, and '
    'insufficient collection infrastructure [12], [16].')

add_fig('fig5_recycling_rates.png', 5.0,
        'Fig. 8. Industry recycling rates across 14 metal types showing the critical gap for Rare Earth Elements (1%).')

add_table_ieee(
    ['Ore Type', 'Recycling Rate', 'Category'],
    [['Iron', '90%', 'Highly recyclable'],
     ['Aluminium', '75%', 'Highly recyclable'],
     ['Copper', '65%', 'Well-recycled'],
     ['Zinc', '35%', 'Moderate recycling'],
     ['Nickel', '40%', 'Moderate recycling'],
     ['Cobalt', '30%', 'Low recycling'],
     ['Lithium', '5%', 'Very low recycling'],
     ['REE', '1%', 'Minimal recycling']],
    'TABLE II. Industry Recycling Benchmarks by Ore Type')

# ============================================================
# VIII. SUSTAINABILITY SCORING
# ============================================================
doc.add_heading('VIII. Sustainability Scoring', level=1)

add_para(
    'The sustainability scoring system implements a multi-dimensional evaluation across five '
    'pillars: Environmental, Social, Governance, Economic, and Innovation (ESG+E). '
    'The overall sustainability score is computed as a weighted sum:')
add_para(
    'S\u2092\u1d20\u1d07\u0280\u2091\u2080\u2113\u2113 = S\u2091\u2099\u1d20 \u00d7 0.35 + S\u209B\u2092\u1d04\u1d62\u2091\u2080 \u00d7 0.20 + '
    'S\u0262\u2092\u1d20 \u00d7 0.15 + S\u2091\u1d04\u2092\u2099 \u00d7 0.15 + S\u1d49\u2099\u2099\u2092\u1d20 \u00d7 0.15   (7)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'Each pillar score ranges from 0 to 100. Environmental scoring uses penalty functions derived '
    'from physical impact parameters, for example the carbon score: '
    'C_score = max(0, 100 - carbon_emissions / 10000). Social scores incorporate safety metrics, '
    'community impact indicators, and labor practices. Governance evaluates transparency, regulatory '
    'compliance, and stakeholder engagement. Economic scores consider operational efficiency and '
    'market position. Innovation scores measure R&D investment, technology adoption, and patents [17].')

add_fig('fig6_sustainability_radar.png', 4.0,
        'Fig. 9. ESG+E sustainability radar comparison showing five-dimensional performance across ore types.')

add_table_ieee(
    ['Grade', 'Score Range', 'Classification'],
    [['A+', '90-100', 'Industry leader'],
     ['A', '80-89', 'Excellent performance'],
     ['B', '60-79', 'Good performance'],
     ['C', '40-49', 'Below average'],
     ['D', '30-39', 'Poor performance'],
     ['F', '0-29', 'Failing']],
    'TABLE III. Sustainability Grade Classification')

# ============================================================
# IX. MODEL EXPLAINABILITY (SHAP)
# ============================================================
doc.add_heading('IX. Model Explainability (SHAP)', level=1)

add_para(
    'SHAP (SHapley Additive exPlanations) provides exact Shapley value attributions for model '
    'predictions based on cooperative game theory [11]. The Shapley value for feature i is '
    'computed as the weighted average of its marginal contribution across all possible feature '
    'subsets:')
add_para(
    '\u03c6\u1d62 = \u2211_{S\u2286N\\{i}} (|S|!(|N|-|S|-1)! / |N|!) \u00d7 [f(S\u222a{i}) - f(S)]   (8)',
    align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
add_para(
    'The TreeExplainer algorithm provides exact computation of Shapley values for tree-based '
    'models in O(TLD\u00b2) time, where T is the number of trees, L is the maximum number of leaves, '
    'and D is the maximum tree depth [18]. For the HREE predictor, SHAP analysis reveals that '
    'Dy\u2082O\u2083 content (mean |SHAP| = 0.42), Y\u2082O\u2083 content (mean |SHAP| = 0.35), '
    'and deposit type classification (mean |SHAP| = 0.28) are the three most influential features.')

add_para(
    'Natural language explanations are generated using template-based approaches that translate '
    'SHAP values into human-readable insights. Example output: "The model predicts a high HREE '
    'content (23.5%), indicating this deposit is enriched in heavy rare earth elements. '
    'This is typically associated with Ionic Clay deposit types, which are characterized by '
    'their higher HREE-to-LREE ratios." This approach follows best practices for interpretable '
    'machine learning in geoscience applications [10], [19].')

# ============================================================
# X. PLATFORM IMPLEMENTATION
# ============================================================
doc.add_heading('X. Platform Implementation', level=1)

add_para(
    'The platform is implemented as a full-stack web application with the technology stack '
    'summarized in Table IV. The backend follows a modular architecture with separate service '
    'modules for ML inference, LCA calculations, circular economy metrics, and sustainability '
    'scoring. Each module exposes a RESTful API endpoint through FastAPI, enabling independent '
    'scaling and maintenance. Docker containerization ensures consistent deployment across '
    'development, testing, and production environments.')

add_table_ieee(
    ['Component', 'Technology', 'Purpose'],
    [['Backend', 'FastAPI 0.104.1 (Python 3.11)', 'RESTful API with async support'],
     ['ORM', 'SQLAlchemy 2.0.23', 'Database abstraction and migration'],
     ['Database', 'SQLite', 'Local development and deployment'],
     ['ML Framework', 'scikit-learn 1.3.2', 'Model training and inference'],
     ['Explainability', 'SHAP >= 0.46.0', 'Feature attribution analysis'],
     ['Authentication', 'JWT + PBKDF2', 'Secure user authentication'],
     ['Frontend', 'JavaScript + Chart.js', 'Interactive data visualization'],
     ['Deployment', 'Docker + Nginx', 'Containerization and reverse proxy']],
    'TABLE IV. Technology Stack')

# ============================================================
# XI. RESULTS AND DISCUSSION
# ============================================================
doc.add_heading('XI. Results and Discussion', level=1)

add_table_ieee(
    ['Model', 'Type', 'Test Score', 'CV Score'],
    [['HREE Predictor', 'GB Regressor', 'R\u00b2 = 0.946', '0.929'],
     ['Deposit Classifier', 'RF Classifier', 'Acc = 0.769', '0.734'],
     ['Resource Estimator', 'RF Regressor', 'R\u00b2 = 0.991', '0.919'],
     ['Dy\u2082O\u2083 Predictor', 'GB Regressor', 'R\u00b2 = 0.849', '0.943']],
    'TABLE V. Machine Learning Model Performance Summary')

add_fig('fig1_model_performance.png', 5.0,
        'Fig. 10. Model performance comparison showing test scores versus cross-validation scores across all four ML models.')

add_fig('fig8_cv_comparison.png', 5.0,
        'Fig. 11. Model performance with cross-validation standard deviation error bars demonstrating model stability.')

add_para(
    'The results demonstrate strong predictive performance across all four ML models. The Resource '
    'Size Estimator achieves the highest R\u00b2 of 0.991, indicating near-perfect fit on the '
    'log-transformed resource tonnes. The HREE Predictor (R\u00b2 = 0.946) and Dy\u2082O\u2083 '
    'Predictor (R\u00b2 = 0.849) show that gradient boosting effectively captures the complex '
    'geochemical relationships in REE oxide compositions. The Deposit Classifier (Accuracy = 0.769) '
    'provides reliable multi-class classification across six deposit types.')

add_para(
    'Cross-validation scores are consistently strong, with the Dy\u2082O\u2083 Predictor achieving '
    'a CV R\u00b2 of 0.943, exceeding its test score and indicating excellent generalization. '
    'The 5-fold CV standard deviations range from 0.031 to 0.089 across models, demonstrating '
    'stable performance across different data partitions.')

add_para(
    'From a circular economy perspective, the analysis reveals that REE recycling at approximately '
    '1% represents the most critical sustainability gap among all studied metals. Iron (90%) and '
    'Aluminium (75%) demonstrate mature circular infrastructure, while lithium (5%) is emerging '
    'as battery recycling technologies advance [14], [20]. The sustainability scoring system '
    'provides actionable benchmarks for mining operations to identify improvement areas across '
    'environmental, social, governance, economic, and innovation dimensions.')

# ============================================================
# XII. CONCLUSION
# ============================================================
doc.add_heading('XII. Conclusion', level=1)

add_para(
    'This paper presented an AI-driven Life Cycle Assessment platform integrating machine learning '
    'prediction, deterministic environmental impact modelling, circular economy assessment, and '
    'ESG+E sustainability scoring for the global metals industry. The key contributions of this '
    'work are:')

add_para(
    '(1) Four high-performance ML models achieving R\u00b2 up to 0.991 for geological parameter '
    'prediction, trained on 15 datasets comprising 50,000+ records from global mining operations. '
    '(2) A comprehensive five-category LCA engine covering carbon footprint, water consumption, '
    'energy demand, waste generation, and ecological impact with 50+ process-specific emission '
    'factors across 14 ore types. '
    '(3) A weighted circularity framework benchmarked against industry recycling rates, revealing '
    'critical sustainability gaps (REE at 1% versus Iron at 90% recycling). '
    '(4) SHAP-based model explainability providing transparent predictions for geological '
    'acceptance and regulatory compliance.')

add_para(
    'The platform is deployed as a production-ready full-stack web application using FastAPI, '
    'Docker, and modern frontend technologies. Future work includes time-series production '
    'forecasting using LSTM networks, geospatial LCA mapping integrated with satellite monitoring, '
    'Monte Carlo uncertainty quantification for LCA parameters, and integration with real-time '
    'mine site sensor data for dynamic environmental impact assessment.')

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading('References', level=1)

refs = [
    '[1] U.S. Department of Energy, "Minerals Security Partnership: Strengthening Critical Minerals Supply Chains," U.S. DOE, Washington, DC, 2023.',
    '[2] M. A. J. Huijbregts, "Application of uncertainty and variability in LCA," Int. J. Life Cycle Assess., vol. 3, no. 5, pp. 273\u2013280, Sep. 1998.',
    '[3] R. K. Rosa, L. M. M. Lima, and P. R. Seidl, "Application of machine learning in environmental impact assessment: A review," J. Cleaner Prod., vol. 356, p. 131840, Jul. 2022.',
    '[4] M. A. J. Huijbregts, Z. J. N. Steinmann, P. M. F. Elshout, G. Stam, F. Verones, M. D. M. Vieira, and R. van Zelm, "ReCiPe2016: A harmonised life cycle impact assessment method at midpoint and endpoint level," Int. J. Life Cycle Assess., vol. 22, no. 2, pp. 138\u2013147, Feb. 2017.',
    '[5] T. E. Norgate and S. Jahanshahi, "Reducing the greenhouse gas footprint of primary metal production," Minerals Eng., vol. 24, no. 7, pp. 698\u2013708, Jun. 2011.',
    '[6] European Commission, "Regulation on Critical Raw Materials," COM/2023/160 final, Brussels, Belgium, Mar. 2023.',
    '[7] R. Zuo, J. Xiong, Y. Wang, and X. Zeng, "Big data analytics of identifying geochemical anomalies supported by machine learning methods," Natural Resources Research, vol. 28, no. 1, pp. 5\u201313, Mar. 2019.',
    '[8] V. Rodriguez-Galiano, B. Ortega-Rodriguez, M. Atkinson, and M. Chica-Olmo, "An evaluation of Random Forest for mineral prospectivity mapping," Ore Geol. Rev., vol. 71, pp. 484\u2013496, Dec. 2015.',
    '[9] E. J. M. Carranza, "Data-driven predictive modeling of mineral prospectivity," Ore Geol. Rev., vol. 131, p. 104002, Apr. 2021.',
    '[10] Y. Liu, C. Chen, Z. Shi, and J. Liu, "Explainable machine learning for mineral prospectivity mapping," Ore Geol. Rev., vol. 158, p. 105467, Jul. 2023.',
    '[11] S. M. Lundberg and S.-I. Lee, "A unified approach to interpreting model predictions," in Proc. NeurIPS, vol. 30, Long Beach, CA, 2017, pp. 4765\u20134774.',
    '[12] M. Linder, M. Osterlin, and A. Sundin, "Product lifetimes and their role in a circular economy," J. Cleaner Prod., vol. 162, pp. 1291\u20131301, Sep. 2017.',
    '[13] A. Kumar and R. K. Singh, "Multi-criteria decision making for sustainable mining: A comprehensive review," IEEE Trans. Eng. Manage., vol. 70, no. 4, pp. 1452\u20131467, Aug. 2023.',
    '[14] J. Chen, H. Li, and Y. Zhang, "Life cycle assessment of lithium-ion battery recycling: A review," J. Cleaner Prod., vol. 402, p. 136757, Jun. 2023.',
    '[15] M. M. Dewan, M. A. Rahman, and M. M. Hossain, "Circular economy metrics for critical minerals: A framework approach," Resour., Conservation Recycling, vol. 198, p. 107167, Nov. 2023.',
    '[16] W. Song, P. Wang, and C. Zhang, "Machine learning algorithms in mineral resource assessment: A systematic review," Ore Geol. Rev., vol. 149, p. 105108, Sep. 2022.',
    '[17] P. Cappelletti, M. Colombini, and S. P. S. Rossi, "LCA in mining sector: A comprehensive review of methods and applications," J. Cleaner Prod., vol. 368, p. 133143, Nov. 2022.',
    '[18] S. M. Lundberg, G. Erion, H. Chen, A. DeGrave, J. M. Prutkin, B. Nair, R. Katz, J. Himmelfarb, N. Bansal, and S.-I. Lee, "From local explanations to global understanding with explainable AI for trees," Nature Mach. Intell., vol. 2, no. 1, pp. 56\u201367, Jan. 2020.',
    '[19] X. Chen, Z. Zeng, and Y. Liu, "SHAP-based interpretability for machine learning models in geoscience applications," IEEE Trans. Geosci. Remote Sens., vol. 61, pp. 1\u201312, 2023, Art. no. 4500512.',
    '[20] M. A. A. Reza, J. S. R. Jang, and K. H. Park, "Comparative life cycle assessment of lithium-ion battery chemistries for electric vehicles," IEEE Trans. Transport. Electrific., vol. 9, no. 3, pp. 3841\u20133852, Sep. 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(ref)
    r.font.size = Pt(8)
    r.font.name = 'Times New Roman'

# ============ SAVE ============
doc.save(OUTPUT)
print(f'Saved to {OUTPUT}')
