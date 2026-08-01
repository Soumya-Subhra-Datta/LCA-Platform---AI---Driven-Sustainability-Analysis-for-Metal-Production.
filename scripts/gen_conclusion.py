from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(3)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_para(text, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=3, space_before=3):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

# Heading
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(9)
p.paragraph_format.space_after = Pt(6)
r = p.add_run('XII. Conclusion')
r.bold = False
r.font.size = Pt(11)
r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0, 0, 0)

add_para(
    'This paper presented an AI-driven Life Cycle Assessment platform that integrates machine learning '
    'prediction, deterministic environmental impact modelling, circular economy assessment, and '
    'ESG+E sustainability scoring for the global metals industry. The platform addresses critical gaps '
    'in conventional LCA methodologies through data-driven automation and cross-commodity comparability.'
)

add_para(
    'Four machine learning models were developed and validated on 15 mining datasets comprising '
    '50,000+ records. The Gradient Boosting Regressor for HREE prediction achieved an R\u00b2 of 0.946, '
    'the Random Forest Classifier for deposit classification reached 76.9% accuracy across six deposit '
    'types, the Random Forest Regressor for resource estimation attained an R\u00b2 of 0.991, and the '
    'Gradient Boosting Regressor for Dy\u2082O\u2083 content prediction achieved an R\u00b2 of 0.849. '
    'These results demonstrate that ensemble tree-based methods, combined with domain-specific feature '
    'engineering, can effectively model complex geochemical and geological relationships.'
)

add_para(
    'The LCA engine models five environmental impact categories\u2014carbon footprint, water consumption, '
    'energy demand, ecological toxicity, and land use\u2014across 14 ore types using ore-specific multipliers '
    'calibrated to industry benchmarks. Analysis reveals that rare earth element production carries the '
    'highest environmental burden, with bastnasite and monazite processing generating 12\u201315 kg CO\u2082e/kg, '
    'compared to 0.3\u20130.4 kg CO\u2082e/kg for iron ore and bauxite.'
)

add_para(
    'The circular economy module quantifies material circularity across five metrics, identifying critical '
    'sustainability gaps: REE recycling rates below 20% contrast sharply with ferrous metals at 85%, '
    'highlighting the urgent need for improved collection and recycling infrastructure. The ESG+E scoring '
    'system provides a multi-dimensional sustainability assessment framework with SHAP-based explainability '
    'ensuring transparent, interpretable predictions for regulatory and investment decision-making.'
)

add_para(
    'The platform is deployed as a production-ready full-stack web application using FastAPI, Docker, '
    'and Chart.js-based frontend visualization. Future work includes time-series production forecasting '
    'using LSTM networks, geospatial LCA mapping integrated with satellite monitoring, Monte Carlo '
    'uncertainty quantification for LCA parameters, and real-time mine site sensor integration for '
    'dynamic environmental impact assessment.'
)

OUTPUT = 'C:/Users/soumy/AppData/Local/Temp/Conclusion.docx'
doc.save(OUTPUT)
print(f'Saved to {OUTPUT}')
